"""Standalone appendix table: GHG-core inputs (combustion proxy, VIIRS flaring)
and the combined core score, for the six worked-example suppliers.

Combines the data behind Figure 7.A (paired bars) and its core companion.
Read-only on the sweep CSV; writes a single-table docx. Does not touch the report.

    python tools/build_ghg_core_table.py
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import docx

import tools.build_report_tables as brt

REPO = Path(__file__).resolve().parents[1]
CSV = REPO / "tools" / "report_example_sweep.csv"
OUT = REPO / "docs" / "ghg_core_inputs_table.docx"

# Weights — CORE_GHG_AUDIT_SUPPORT_WEIGHTS (engine/constants.py:330).
W_COMB, W_FLARE = 0.60, 0.40
SITE_TYPE = {
    "Jamnagar": "refinery", "Comodoro": "oil & gas field", "Norilsk": "smelter",
    "Morowali": "nickel smelter", "Escondida": "copper mine",
    "Carajás": "iron ore mine",
}


def main():
    S = {r["site"]: r for r in csv.DictReader(CSV.open(encoding="utf-8"))}

    def sk(n):
        return next(k for k in S if n.lower() in k.lower())

    rows = []
    for needle in SITE_TYPE:
        r = S[sk(needle)]
        comb = float(r["ghg_combustion_proxy"])
        flare = float(r["ghg_viirs_flaring_score"])
        core = W_COMB * comb + W_FLARE * flare
        rows.append((needle, SITE_TYPE[needle], comb, flare, core))
    rows.sort(key=lambda t: -t[4])  # by GHG core, descending

    doc = docx.Document()
    headers = ["Site", "Site type", "Combustion proxy", "VIIRS flaring",
               "GHG core"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    right = {2, 3, 4}
    for c, h in enumerate(headers):
        brt._set_cell(table.rows[0].cells[c], h, bold=True, fill=brt.HEADER_FILL,
                      align="right" if c in right else None)
    for i, (name, typ, comb, flare, core) in enumerate(rows):
        cells = [name, typ, f"{comb:.3f}", f"{flare:.3f}", f"{core:.3f}"]
        for c, val in enumerate(cells):
            brt._set_cell(table.rows[i + 1].cells[c], val,
                          bold=(c == 4),  # emphasise the combined core
                          align="right" if c in right else None)

    cap = doc.add_paragraph(style="Caption")
    cap.add_run(
        "Table X.?: GHG-core inputs and output for the six worked-example "
        "suppliers — the data behind Figure 7.A and its companion combined. "
        "The combustion proxy (NO₂ + CO, borrowed from the Air pillar) and the "
        "VIIRS flaring score are the two scored inputs; the GHG core is their "
        "weighted combination, 0.60 × combustion + 0.40 × flaring. All values "
        "are dimensionless on [0, 1]; rows ordered by GHG core. Source: "
        "Appendix X sweep."
    ).italic = True

    doc.save(str(OUT))
    print(f"wrote {OUT}\n")
    print("| Site | Site type | Combustion proxy | VIIRS flaring | GHG core |")
    print("|---|---|---|---|---|")
    for name, typ, comb, flare, core in rows:
        print(f"| {name} | {typ} | {comb:.3f} | {flare:.3f} | **{core:.3f}** |")


if __name__ == "__main__":
    main()
