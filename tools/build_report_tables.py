"""Insert the Section 7 worked-example tables into the GSCO report docx.

Reads the sweep outputs (tools/report_example_sweep.csv summary +
tools/report_example_perindicator.csv long-format), builds python-docx native
tables (style "Table Grid", bold + light-grey-shaded header, numeric columns
right-aligned), inserts each immediately BEFORE its "[Table 7.x: …]" placeholder
paragraph with an italic caption below, then deletes the placeholder.

Saves to a NEW filename — never overwrites the source draft.

Read-only on engine/constants/docs; this script only edits a COPY of the docx.

    python tools/build_report_tables.py
"""

from __future__ import annotations

import csv
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPO = Path(__file__).resolve().parents[1]
SRC_DOCX = REPO / "docs" / "GSCO_Report_Draft2.docx"
OUT_DOCX = REPO / "docs" / "GSCO_Report_Draft2_Section7_tables.docx"
SUMMARY_CSV = REPO / "tools" / "report_example_sweep.csv"
LONG_CSV = REPO / "tools" / "report_example_perindicator.csv"

HEADER_FILL = "D9D9D9"   # light grey
DASH = "—"


# ---------------------------------------------------------------------------
# CSV loading
# ---------------------------------------------------------------------------

def _load_summary() -> dict[str, dict]:
    with SUMMARY_CSV.open(encoding="utf-8") as fh:
        return {r["site"]: r for r in csv.DictReader(fh)}


def _load_long() -> dict[tuple[str, str], dict]:
    out: dict[tuple[str, str], dict] = {}
    with LONG_CSV.open(encoding="utf-8") as fh:
        for r in csv.DictReader(fh):
            out[(r["site"], r["indicator"])] = r
    return out


def _site_key(summary: dict, needle: str) -> str:
    """Resolve a full CSV site key from a substring (e.g. 'Escondida')."""
    for name in summary:
        if needle.lower() in name.lower():
            return name
    raise KeyError(f"no site matches {needle!r}")


# ---------------------------------------------------------------------------
# Number formatting
# ---------------------------------------------------------------------------

def _to_float(v):
    if v is None or v == "" or v == "None":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def score3(v) -> str:
    f = _to_float(v)
    return DASH if f is None else f"{f:.3f}"


def ha0(v) -> str:
    f = _to_float(v)
    return DASH if f is None else f"{round(f):,}"


def pct1(v) -> str:
    f = _to_float(v)
    return DASH if f is None else f"{f:.1f}"


def num(v, dp=3) -> str:
    f = _to_float(v)
    return DASH if f is None else f"{f:.{dp}f}"


# ---------------------------------------------------------------------------
# docx low-level helpers
# ---------------------------------------------------------------------------

def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell(cell, text: str, *, bold=False, italic=False,
              align: str | None = None, fill: str | None = None) -> None:
    # Fresh add_table cells have one empty paragraph with no runs — add a
    # single run so the only run carries the bold/italic formatting.
    para = cell.paragraphs[0]
    run = para.add_run("" if text is None else str(text))
    run.bold = bold
    run.italic = italic
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fill:
        _shade(cell, fill)


def _find_placeholder(doc, label: str):
    """Return the paragraph object whose text contains '<label>:' (e.g. 'Table 7.4')."""
    target = f"{label}:"
    for p in doc.paragraphs:
        if target in p.text:
            return p
    raise KeyError(f"placeholder for {label!r} not found")


def _insert_table_and_caption(doc, placeholder, headers, rows,
                              right_cols: set[int], caption: str,
                              bold_cells: set[tuple[int, int]] | None = None,
                              delete_placeholder: bool = True):
    """Build a Table-Grid table + italic caption, insert both before the
    placeholder paragraph, then (optionally) delete the placeholder.

    `rows` is a list of row-tuples of display strings. `right_cols` are the
    0-based column indices to right-align in the body. `bold_cells` is a set of
    (row_index, col_index) data-cell coordinates to bold (row_index is 0-based
    over `rows`). Set `delete_placeholder=False` to stack a second table before
    the same placeholder (used for 7.7 + 7.7b).
    """
    bold_cells = bold_cells or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header.
    for c, htext in enumerate(headers):
        _set_cell(table.rows[0].cells[c], htext, bold=True, fill=HEADER_FILL,
                  align="right" if c in right_cols else None)
    # Body.
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            _set_cell(
                table.rows[r + 1].cells[c], val,
                bold=(r, c) in bold_cells,
                align="right" if c in right_cols else None,
            )

    # Move table immediately before the placeholder.
    placeholder._p.addprevious(table._tbl)
    # Caption (italic) immediately below the table, above the placeholder.
    cap = doc.add_paragraph(style="Caption")
    crun = cap.add_run(caption)
    crun.italic = True
    placeholder._p.addprevious(cap._p)
    # Remove the placeholder (unless stacking another table before it).
    if delete_placeholder:
        placeholder._p.getparent().remove(placeholder._p)


# ---------------------------------------------------------------------------
# Table builders — return (status, dash_note)
# ---------------------------------------------------------------------------

# Short display labels for narrow cells.
SHORT = {
    "Morowali Industrial Park (IMIP)": "Morowali IMIP",
    "CATL battery plant, Ningde": "CATL Ningde",
    "Carajás iron ore mine": "Carajás",
    "Jamnagar refinery": "Jamnagar",
    "Escondida copper mine": "Escondida",
    "Norilsk Nickel — Nadezhda smelter": "Norilsk",
    "Comodoro Rivadavia oil & gas": "Comodoro (Patagonia)",
    "Tongwei/LONGi polysilicon, Leshan": "Leshan (LONGi)",
}


def _short(name: str) -> str:
    return SHORT.get(name, name)


def build_70(doc, summary):
    """Headline outputs — transposed (metrics as rows, two sites as columns)."""
    a = summary[_site_key(summary, "Morowali")]
    b = summary[_site_key(summary, "CATL")]
    metrics = [
        ("Composite screening", "composite_overall", score3),
        ("Traffic-light band", "composite_band", str),
        ("Composite confidence", "composite_confidence", score3),
        ("Air follow-up", "air_followup", score3),
        ("GHG follow-up", "ghg_followup", score3),
        ("Nature follow-up", "nature_followup", score3),
    ]
    headers = ["Metric", _short(a["site"]), _short(b["site"])]
    rows = [(label, fmt(a[key]), fmt(b[key])) for label, key, fmt in metrics]
    cap = (f"Table 7.0: Headline screening outputs — {_short(a['site'])} "
           f"(high-signal, Moderate band) vs {_short(b['site'])} (Low-band "
           f"baseline). Composite = mean of the three pillar follow-ups.")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.0"),
                              headers, rows, {1, 2}, cap)
    return "full", None


def build_71(doc, long):
    """Repeatable-core trace on one clean Air indicator (Escondida NO₂)."""
    site = _site_key({r[0]: 1 for r in long}, "Escondida")
    rec = long[(site, "air.no2")]
    steps = [
        ("Raw site value (NO₂, ring units)", num(rec["raw_value"], 3)),
        ("Background median (ring)", num(rec["bg_median"], 3)),
        ("Anomaly (site − median)", num(rec["anomaly"], 3)),
        ("Background σ (standardising)", num(rec["bg_std"], 3)),
        ("z-score (anomaly ÷ σ)", num(rec["z"], 3)),
        ("Normalised score (0–1)", score3(rec["score"])),
    ]
    headers = ["Step", "Value"]
    cap = ("Table 7.1: The repeatable-core anomaly trace on a single indicator "
           "(Escondida, NO₂). Each step is verifiable: site − median = anomaly; "
           "anomaly ÷ σ = z; z maps to the normalised score. σ is the "
           "standardising (climatology-baseline) spread that reconciles z.")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.1"),
                              headers, steps, {1}, cap)
    return "full", None


def build_73(doc, long):
    """Per-indicator Air z/score for Escondida vs Comodoro (PM omitted)."""
    sa = _site_key({r[0]: 1 for r in long}, "Escondida")
    sb = _site_key({r[0]: 1 for r in long}, "Comodoro")
    pollutants = [
        ("NO₂", "air.no2"), ("SO₂", "air.so2"), ("CO", "air.co"),
        ("HCHO", "air.hcho"), ("O₃", "air.o3"), ("AAI", "air.aai"),
        ("AOD", "air.aod"),
    ]
    headers = ["Indicator", "Esc. z", "Esc. score", "Com. z", "Com. score"]
    rows = []
    n_dash = 0
    for label, ind in pollutants:
        ra = long.get((sa, ind), {})
        rb = long.get((sb, ind), {})
        cells = (label, num(ra.get("z"), 3), score3(ra.get("score")),
                 num(rb.get("z"), 3), score3(rb.get("score")))
        n_dash += sum(1 for c in cells[1:] if c == DASH)
        rows.append(cells)
    cap = ("Table 7.3: Per-indicator Air results (z and normalised score) — "
           "Escondida vs Comodoro Rivadavia. PM₂.₅ and PM₁₀ are omitted: both "
           "dropped because the 5 km site buffer is smaller than their 44.5 km "
           "native pixel. Cells shown as — had no valid retrieval in-window.")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.3"),
                              headers, rows, {1, 2, 3, 4}, cap)
    return ("full" if n_dash == 0 else "partial",
            f"{n_dash} — cells (no valid retrieval)" if n_dash else None)


def build_74(doc, summary):
    """Scored GHG core & follow-up — Jamnagar vs Escondida."""
    a = summary[_site_key(summary, "Jamnagar")]
    b = summary[_site_key(summary, "Escondida")]
    headers = ["Site", "Combustion proxy", "VIIRS flaring", "GHG follow-up"]
    rows = [
        (_short(a["site"]), score3(a["ghg_combustion_proxy"]),
         score3(a["ghg_viirs_flaring_score"]), score3(a["ghg_followup"])),
        (_short(b["site"]), score3(b["ghg_combustion_proxy"]),
         score3(b["ghg_viirs_flaring_score"]), score3(b["ghg_followup"])),
    ]
    cap = ("Table 7.4: Scored GHG core and follow-up — Jamnagar (combustion "
           "proxy ≈ 0.05 but VIIRS flaring = 1.0: the two grammars diverge, "
           "flaring catches what the Air borrow misses) vs Escondida "
           "(combustion ≈ VIIRS: the grammars converge).")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.4"),
                              headers, rows, {1, 2, 3}, cap)
    return "full", None


def build_75(doc, long):
    """GHG reference/context raw values — Jamnagar vs Escondida (not scored)."""
    sa = _site_key({r[0]: 1 for r in long}, "Jamnagar")
    sb = _site_key({r[0]: 1 for r in long}, "Escondida")

    def _co2_ri(site):
        note = long.get((site, "ghg.co2"), {}).get("extra_note", "")
        # note: "context — not scored; rel_intensity=10.0 (cap 10), total..."
        for tok in note.split(";"):
            if "rel_intensity" in tok:
                return tok.split("rel_intensity=")[1].split(" ")[0]
        return None

    headers = ["Site", "CO₂ raw mean", "CO₂ rel. int.", "VIIRS bright.",
               "CH₄ z", "CH₄ score"]
    rows = []
    for site in (sa, sb):
        co2 = long.get((site, "ghg.co2"), {})
        viirs = long.get((site, "ghg.viirs"), {})
        ch4 = long.get((site, "ghg.ch4"), {})
        rows.append((
            _short(site), num(co2.get("raw_value"), 1), num(_co2_ri(site), 1),
            num(viirs.get("raw_value"), 3), num(ch4.get("z"), 3),
            score3(ch4.get("score")),
        ))
    cap = ("Table 7.5: GHG reference/context terms (displayed, NOT scored) — "
           "Jamnagar vs Escondida. CO₂ (ODIAC inventory) relative intensity is "
           "capped at 10; CH₄ is an M-CH4-A1 reference signal. None of these "
           "enter the GHG core, which is combustion + VIIRS flaring.")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.5"),
                              headers, rows, {1, 2, 3, 4, 5}, cap)
    return "full", None


def build_76(doc, summary):
    """Per-strand Nature — Carajás vs Morowali."""
    a = summary[_site_key(summary, "Carajás")]
    b = summary[_site_key(summary, "Morowali")]
    headers = ["Site", "KBA overlap (ha)", "KBA dist (km)",
               "Proximity branch", "Habitat conv.", "Nature follow-up"]
    rows = []
    for s in (a, b):
        overlap = _to_float(s["nature_kba_overlap_ha"])
        branch = "overlap" if (overlap or 0) > 0 else "distance-decay"
        rows.append((
            _short(s["site"]), ha0(s["nature_kba_overlap_ha"]),
            num(s["nature_kba_dist_km"], 2), branch,
            score3(s["nature_habitat_conversion"]),
            score3(s["nature_followup"]),
        ))
    cap = ("Table 7.6: Per-strand Nature results — Carajás (sits inside a KBA: "
           "98.8 % overlap → proximity via the overlap branch) vs Morowali IMIP "
           "(KBA ~6 km away, 0 overlap → proximity via the distance-decay "
           "branch; highest on-site habitat conversion). Proximity score = "
           "max(overlap_pct, exp(−dist/10)).")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.6"),
                              headers, rows, {1, 2, 4, 5}, cap)
    return "full", None


def build_77(doc, summary, long):
    """Pillar-level confidence + binding pillar (7.7), then the confidence
    sub-terms companion (7.7b) stacked immediately after, both before the same
    7.7 placeholder."""
    placeholder = _find_placeholder(doc, "Table 7.7")
    a = summary[_site_key(summary, "Norilsk")]
    b = summary[_site_key(summary, "Escondida")]

    # --- 7.7 main ---
    headers = ["Site", "Air conf", "GHG conf", "Nature conf",
               "Composite (min)", "Binding pillar"]
    rows, bold = [], set()
    for ri, s in enumerate((a, b)):
        confs = {
            "Air": _to_float(s["air_confidence"]),
            "GHG": _to_float(s["ghg_confidence"]),
            "Nature": _to_float(s["nature_confidence"]),
        }
        present = {k: v for k, v in confs.items() if v is not None}
        binding = min(present, key=present.get) if present else DASH
        rows.append((
            _short(s["site"]), score3(s["air_confidence"]),
            score3(s["ghg_confidence"]), score3(s["nature_confidence"]),
            score3(s["composite_confidence"]), binding,
        ))
        bold.add((ri, 4))   # composite (min) cell
        bold.add((ri, 5))   # binding pillar name
    cap = ("Table 7.7: Pillar-level confidence and the binding (worst) pillar — "
           "Norilsk vs Escondida. Composite confidence is the per-pillar "
           "minimum (bold); the binding pillar (bold) is the one that sets it.")
    _insert_table_and_caption(doc, placeholder, headers, rows, {1, 2, 3, 4},
                              cap, bold_cells=bold, delete_placeholder=False)

    # --- 7.7b companion (sub-terms) ---
    sa = _site_key({r[0]: 1 for r in long}, "Norilsk")
    sb = _site_key({r[0]: 1 for r in long}, "Escondida")
    reps = [("Air — NO₂", "air.no2"), ("GHG — VIIRS", "ghg.viirs")]
    headers_b = ["Site / indicator", "QA", "N_valid", "Anomaly str.", "Spatial ctx"]
    rows_b = []
    n_dash = 0
    for site in (sa, sb):
        for label, ind in reps:
            rec = long.get((site, ind), {})
            cells = (f"{_short(site)} · {label}", num(rec.get("qa"), 2),
                     num(rec.get("n_valid"), 2), num(rec.get("anomaly_strength"), 3),
                     num(rec.get("spatial_context"), 2))
            n_dash += sum(1 for c in cells[1:] if c == DASH)
            rows_b.append(cells)
    rows_b.append(("Nature strands (KBA/habitat/NDVI)", DASH, DASH, DASH, DASH))
    cap_b = ("Table 7.7b: Confidence sub-terms for representative indicators "
             "(Air = NO₂, GHG = VIIRS) at the two worked sites. Nature strands "
             "(KBA, habitat, NDVI) use a coverage/cloud-based confidence that is "
             "not decomposed into these four terms — shown as —.")
    _insert_table_and_caption(doc, placeholder, headers_b, rows_b,
                              {1, 2, 3, 4}, cap_b, delete_placeholder=True)
    return ("partial", "7.7b Nature sub-terms — (separate confidence construction)")


def build_710(doc, summary):
    """Attributability labels — 3 flaring sites + Leshan."""
    order = ["Morowali", "Comodoro", "Jamnagar", "Leshan"]
    headers = ["Site", "VIIRS lit-contrast (state · pct)",
               "Habitat spatial-link (state · offset km)", "Wind states"]
    rows = []
    n_dash = 0
    for needle in order:
        s = summary[_site_key(summary, needle)]
        viirs = (f"{s['ghg_viirs_attrib_state'] or DASH} · "
                 f"{pct1(_to_float(s['ghg_viirs_lit_contrast_pct']) * 100 if _to_float(s['ghg_viirs_lit_contrast_pct']) is not None else None)}%")
        off = s["nature_spatial_offset_km"]
        hab = (f"{s['nature_habitat_attrib_state'] or DASH} · "
               f"{num(off, 2) if _to_float(off) is not None else DASH} km")
        wind = (s["wind_attrib_states"] or DASH).replace(";", ", ")
        if _to_float(off) is None:
            n_dash += 1
        rows.append((_short(s["site"]), viirs, hab, wind))
    cap = ("Table 7.10: Attributability labels at the flaring sites and the "
           "polysilicon site — VIIRS lit-contrast state (percentile within the "
           "ring), habitat spatial-link state (change-centroid offset from the "
           "supplier), and per-pollutant wind states. 'sparse' offsets show as "
           "— (too few change pixels to localise).")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.10"),
                              headers, rows, set(), cap)
    return ("full" if n_dash == 0 else "partial",
            f"{n_dash} habitat-offset cells — (sparse)" if n_dash else None)


def build_711(doc, summary):
    """Degradation events — Norilsk vs Escondida."""
    a = summary[_site_key(summary, "Norilsk")]
    b = summary[_site_key(summary, "Escondida")]
    headers = ["Site", "Dropped indicators", "Fallbacks fired",
               "Survivors (Air)", "Composite conf."]
    rows = []
    for s in (a, b):
        dropped = (s["dropped_indicators"] or "").split(";") if s["dropped_indicators"] else []
        dropped_disp = ", ".join(d.replace(".score", "").replace("air.", "")
                                 for d in dropped) or "none"
        fb = (s["fallbacks_fired"] or "").replace(";", ", ") or "none"
        # Air survivors out of 9 (no2,so2,co,hcho,pm25,pm10,o3,aai,aod).
        air_dropped = sum(1 for d in dropped if d.startswith("air."))
        survivors = f"{9 - air_dropped} of 9"
        rows.append((_short(s["site"]), dropped_disp, fb, survivors,
                     score3(s["composite_confidence"])))
    cap = ("Table 7.11: Degradation events — Norilsk vs Escondida. PM₂.₅ and "
           "PM₁₀ are dropped at every site (5 km buffer < 44.5 km native "
           "pixel), renormalising Air to 7 of 9 survivors. Norilsk additionally "
           "recovered SO₂ and AOD via the SPPY temporal fallback (polar-winter "
           "sparsity); Escondida fired no fallback.")
    _insert_table_and_caption(doc, _find_placeholder(doc, "Table 7.11"),
                              headers, rows, {3, 4}, cap)
    return "full", None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    summary = _load_summary()
    long = _load_long()
    doc = docx.Document(str(SRC_DOCX))

    results: dict[str, tuple[str, str | None]] = {}
    # Each builder locates its own placeholder by text, so paragraph-index
    # shifts from prior inserts don't matter. 7.7 + 7.7b are stacked by a single
    # builder before the shared 7.7 placeholder.
    builders = [
        ("Table 7.0", lambda: build_70(doc, summary)),
        ("Table 7.1", lambda: build_71(doc, long)),
        ("Table 7.3", lambda: build_73(doc, long)),
        ("Table 7.4", lambda: build_74(doc, summary)),
        ("Table 7.5", lambda: build_75(doc, long)),
        ("Table 7.6", lambda: build_76(doc, summary)),
        ("Table 7.7", lambda: build_77(doc, summary, long)),
        ("Table 7.11", lambda: build_711(doc, summary)),
        ("Table 7.10", lambda: build_710(doc, summary)),
    ]
    for label, fn in builders:
        try:
            status, note = fn()
            results[label] = (status, note)
            print(f"[tables] {label}: {status}" + (f"  ({note})" if note else ""))
        except Exception as exc:
            results[label] = ("ERROR", f"{type(exc).__name__}: {exc}")
            print(f"[tables] {label}: ERROR — {exc}")
            raise

    doc.save(str(OUT_DOCX))
    print(f"\n[tables] saved {OUT_DOCX}")
    print("\n=== population summary ===")
    for label, (status, note) in results.items():
        print(f"  {label:12} {status:8} {note or ''}")


if __name__ == "__main__":
    main()
