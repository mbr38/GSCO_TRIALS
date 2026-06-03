"""Standalone table behind Figure 7.B (pillar signatures): the three pillar
follow-up priorities, the composite score, and composite confidence for the two
contrasting worked-example suppliers.

Read-only on the sweep CSV; writes a single-table docx. Does not touch the report.

    python tools/build_pillar_signatures_table.py
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import docx

import tools.build_report_tables as brt

REPO = Path(__file__).resolve().parents[1]
CSV = REPO / "tools" / "report_example_sweep.csv"
OUT = REPO / "docs" / "pillar_signatures_table.docx"

SITES = [("Comodoro", "Comodoro Rivadavia (oil & gas)"),
         ("Carajás", "Carajás (iron ore)")]


def main():
    S = {r["site"]: r for r in csv.DictReader(CSV.open(encoding="utf-8"))}

    def sk(n):
        return next(k for k in S if n.lower() in k.lower())

    doc = docx.Document()
    headers = ["Site", "Air", "GHG", "Nature", "Composite", "Composite confidence"]
    table = doc.add_table(rows=1 + len(SITES), cols=len(headers))
    table.style = "Table Grid"
    right = {1, 2, 3, 4, 5}
    for c, h in enumerate(headers):
        brt._set_cell(table.rows[0].cells[c], h, bold=True, fill=brt.HEADER_FILL,
                      align="right" if c in right else None)
    for i, (needle, label) in enumerate(SITES):
        r = S[sk(needle)]
        cells = [label,
                 f"{float(r['air_followup']):.3f}",
                 f"{float(r['ghg_followup']):.3f}",
                 f"{float(r['nature_followup']):.3f}",
                 f"{float(r['composite_overall']):.3f}",
                 f"{float(r['composite_confidence']):.3f}"]
        for c, val in enumerate(cells):
            brt._set_cell(table.rows[i + 1].cells[c], val,
                          bold=(c == 4),  # emphasise the composite, as in Fig 7.B
                          align="right" if c in right else None)

    cap = doc.add_paragraph(style="Caption")
    cap.add_run(
        "Table X.?: Pillar follow-up priorities, composite screening score and "
        "composite confidence for the two contrasting suppliers in Figure 7.B. "
        "The Air, GHG and Nature columns are the per-pillar follow-up priorities; "
        "the composite screening score is their equal-weighted mean, and "
        "composite confidence is the minimum of the three pillar confidences. "
        "All values are dimensionless on [0, 1]. Source: Appendix X sweep."
    ).italic = True

    doc.save(str(OUT))
    print(f"wrote {OUT}\n")
    print("| Site | Air | GHG | Nature | Composite | Composite confidence |")
    print("|---|---|---|---|---|---|")
    for needle, label in SITES:
        r = S[sk(needle)]
        print(f"| {label} | {float(r['air_followup']):.3f} | "
              f"{float(r['ghg_followup']):.3f} | {float(r['nature_followup']):.3f} | "
              f"**{float(r['composite_overall']):.3f}** | "
              f"{float(r['composite_confidence']):.3f} |")


if __name__ == "__main__":
    main()
