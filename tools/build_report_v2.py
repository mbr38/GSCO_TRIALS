"""Revise Section 7 of the GSCO report: two inline figures, a merged inline
confidence/degradation table, and a new Appendix X holding the per-indicator
detail tables.

Starts from the clean placeholder anchors in docs/GSCO_Report_Draft2.docx
(deterministic — identical body prose to the prior table-insertion pass) and
writes a NEW file. Read-only on engine/constants/docs; figures + CSVs are the
only generated artefacts.

    python tools/report_figures.py        # first: render the two PNGs
    python tools/build_report_v2.py       # then: assemble the docx

Placement decisions (the brief's "Fig 7.B replaces Table 7.0, at §7.5" is
internally inconsistent — the Table 7.0 placeholder sits at §7.0, not §7.5):
  - Figure 7.A  -> at the Table 7.4 placeholder (§7.3 GHG pillar). Placeholder removed.
  - Figure 7.B  -> at §7.5 "Combining pillars: the composite" (before the §7.6
                   heading). The Table 7.0 placeholder is replaced by a one-line
                   cross-reference to Figure 7.B.
  - Table 7.7   -> merged confidence+degradation inline table at its placeholder.
  - 7.1/7.3/7.5/7.6/7.10/7.11 placeholders -> one-line cross-refs to Appendix X.
  - Table 7.2 placeholder -> left untouched (out of scope).
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm

import tools.build_report_tables as brt

REPO = Path(__file__).resolve().parents[1]
SRC_DOCX = REPO / "docs" / "GSCO_Report_Draft2.docx"
OUT_DOCX = REPO / "docs" / "GSCO_Report_Draft2_Section7_v2.docx"
FIG_A = REPO / "tools" / "fig_7a_ghg_slope.png"
FIG_B = REPO / "tools" / "fig_7b_pillar_signatures.png"

DASH = brt.DASH
_to_float = brt._to_float
score3, num, ha0, pct1 = brt.score3, brt.num, brt.ha0, brt.pct1
_set_cell = brt._set_cell
_short = brt._short
_site_key = brt._site_key

# Pretty pollutant labels for the merged table / appendix degradation table.
_PRETTY = {
    "pm25": "PM₂.₅", "pm10": "PM₁₀", "so2": "SO₂",
    "no2": "NO₂", "aod": "AOD", "aai": "AAI", "hcho": "HCHO",
    "co": "CO", "o3": "O₃", "ch4": "CH₄", "co2": "CO₂",
    "viirs": "VIIRS",
}


def _strip_token(tok: str) -> str:
    base = tok.strip()
    for pre in ("air.", "ghg.", "nature."):
        if base.startswith(pre):
            base = base[len(pre):]
    base = base.replace(".score", "")
    return _PRETTY.get(base, base.upper())


def _pretty_list(value: str, suffix: str = "") -> str:
    if not value:
        return "none"
    parts = [p for p in value.split(";") if p]
    labels = [_strip_token(p) for p in parts]
    return (", ".join(labels) + suffix) if labels else "none"


# ---------------------------------------------------------------------------
# docx insertion helpers
# ---------------------------------------------------------------------------

def _caption(doc, text):
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(text).italic = True
    return cap


def _insert_figure_before(doc, anchor, png, caption_text):
    pimg = doc.add_paragraph()
    pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pimg.add_run().add_picture(str(png), width=Cm(15))
    anchor._p.addprevious(pimg._p)
    cap = _caption(doc, caption_text)
    anchor._p.addprevious(cap._p)


def _replace_placeholder_with_figure(doc, label, png, caption_text):
    ph = brt._find_placeholder(doc, label)
    _insert_figure_before(doc, ph, png, caption_text)
    ph._p.getparent().remove(ph._p)


def _crossref(doc, label, text):
    ph = brt._find_placeholder(doc, label)
    for r in list(ph.runs):
        r._r.getparent().remove(r._r)
    ph.style = doc.styles["Normal"]
    ph.add_run(text)


def _find_heading(doc, needle):
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and needle.lower() in p.text.lower():
            return p
    raise KeyError(f"heading containing {needle!r} not found")


def _append_table(doc, headers, rows, right_cols, caption_text,
                  bold_cells=None):
    bold_cells = bold_cells or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        _set_cell(table.rows[0].cells[c], h, bold=True, fill=brt.HEADER_FILL,
                  align="right" if c in right_cols else None)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            _set_cell(table.rows[r + 1].cells[c], val,
                      bold=(r, c) in bold_cells,
                      align="right" if c in right_cols else None)
    _caption(doc, caption_text)
    return table


# ---------------------------------------------------------------------------
# Inline merged Table 7.7 (confidence + degradation)
# ---------------------------------------------------------------------------

def insert_merged_77(doc, summary):
    a = summary[_site_key(summary, "Norilsk")]
    b = summary[_site_key(summary, "Escondida")]
    headers = ["Site", "Air conf", "GHG conf", "Nature conf", "Composite (min)",
               "Dropped indicators", "Fallbacks fired"]
    rows, bold = [], set()
    for ri, s in enumerate((a, b)):
        rows.append((
            _short(s["site"]), score3(s["air_confidence"]),
            score3(s["ghg_confidence"]), score3(s["nature_confidence"]),
            score3(s["composite_confidence"]),
            _pretty_list(s["dropped_indicators"]),
            _pretty_list(s["fallbacks_fired"], suffix=" (SPPY)"),
        ))
        bold.add((ri, 4))
    cap = ("Table 7.7: Per-pillar confidence with the binding composite minimum "
           "(bold), shown with each site's degradation events. PM₂.₅ and "
           "PM₁₀ drop at every screened site (5 km buffer < 44.5 km "
           "native pixel); Norilsk is the only site to fire the SPPY temporal "
           "fallback, recovering SO₂ and AOD from polar-winter sparsity.")
    ph = brt._find_placeholder(doc, "Table 7.7")
    brt._insert_table_and_caption(doc, ph, headers, rows, {1, 2, 3, 4}, cap,
                                  bold_cells=bold)


# ---------------------------------------------------------------------------
# Appendix X data builders (reuse the proven row logic, appendix captions)
# ---------------------------------------------------------------------------

def _long_site(long, needle):
    return _site_key({k[0]: 1 for k in long}, needle)


def appendix_x1(doc, long):
    site = _long_site(long, "Escondida")
    rec = long[(site, "air.no2")]
    rows = [
        ("Raw site value (NO₂, ring units)", num(rec["raw_value"], 3)),
        ("Background median (ring)", num(rec["bg_median"], 3)),
        ("Anomaly (site − median)", num(rec["anomaly"], 3)),
        ("Background σ (standardising)", num(rec["bg_std"], 3)),
        ("z-score (anomaly ÷ σ)", num(rec["z"], 3)),
        ("Normalised score (0–1)", score3(rec["score"])),
    ]
    cap = ("Table X.1: The repeatable-core anomaly trace on a single indicator "
           "(Escondida, NO₂). Each step is verifiable: site − median = "
           "anomaly; anomaly ÷ σ = z; z maps to the normalised score. "
           "σ is the standardising (climatology-baseline) spread.")
    _append_table(doc, ["Step", "Value"], rows, {1}, cap)
    return "full", None


def appendix_x2(doc, long):
    sa = _long_site(long, "Escondida")
    sb = _long_site(long, "Comodoro")
    pollutants = [("NO₂", "air.no2"), ("SO₂", "air.so2"),
                  ("CO", "air.co"), ("HCHO", "air.hcho"), ("O₃", "air.o3"),
                  ("AAI", "air.aai"), ("AOD", "air.aod")]
    headers = ["Indicator", "Esc. z", "Esc. score", "Com. z", "Com. score"]
    rows, n_dash = [], 0
    for label, ind in pollutants:
        ra = long.get((sa, ind), {})
        rb = long.get((sb, ind), {})
        cells = (label, num(ra.get("z"), 3), score3(ra.get("score")),
                 num(rb.get("z"), 3), score3(rb.get("score")))
        n_dash += sum(1 for c in cells[1:] if c == DASH)
        rows.append(cells)
    cap = ("Table X.2: Per-indicator Air results (z and normalised score) — "
           "Escondida vs Comodoro Rivadavia. PM₂.₅ and PM₁₀ "
           "are omitted (dropped: 5 km buffer < 44.5 km native pixel). Cells "
           "shown as — had no valid retrieval in-window.")
    _append_table(doc, headers, rows, {1, 2, 3, 4}, cap)
    return ("full" if n_dash == 0 else "partial",
            f"{n_dash} — cells" if n_dash else None)


def appendix_x3(doc, long):
    sa = _long_site(long, "Jamnagar")
    sb = _long_site(long, "Escondida")

    def _ri(site):
        note = long.get((site, "ghg.co2"), {}).get("extra_note", "")
        for tok in note.split(";"):
            if "rel_intensity" in tok:
                return tok.split("rel_intensity=")[1].split(" ")[0]
        return None

    headers = ["Site", "CO₂ raw mean", "CO₂ rel. int.", "VIIRS bright.",
               "CH₄ z", "CH₄ score"]
    rows = []
    for site in (sa, sb):
        co2 = long.get((site, "ghg.co2"), {})
        viirs = long.get((site, "ghg.viirs"), {})
        ch4 = long.get((site, "ghg.ch4"), {})
        rows.append((_short(site), num(co2.get("raw_value"), 1),
                     num(_ri(site), 1), num(viirs.get("raw_value"), 3),
                     num(ch4.get("z"), 3), score3(ch4.get("score"))))
    cap = ("Table X.3: GHG reference/context terms (displayed, NOT scored) — "
           "Jamnagar vs Escondida. CO₂ (ODIAC) relative intensity is capped "
           "at 10; CH₄ is an M-CH4-A1 reference signal. None enter the GHG "
           "core (combustion + VIIRS flaring).")
    _append_table(doc, headers, rows, {1, 2, 3, 4, 5}, cap)
    return "full", None


def appendix_x4(doc, summary):
    a = summary[_site_key(summary, "Carajás")]
    b = summary[_site_key(summary, "Morowali")]
    headers = ["Site", "KBA overlap (ha)", "KBA dist (km)", "Proximity branch",
               "Habitat conv.", "Nature follow-up"]
    rows = []
    for s in (a, b):
        overlap = _to_float(s["nature_kba_overlap_ha"])
        branch = "overlap" if (overlap or 0) > 0 else "distance-decay"
        rows.append((_short(s["site"]), ha0(s["nature_kba_overlap_ha"]),
                     num(s["nature_kba_dist_km"], 2), branch,
                     score3(s["nature_habitat_conversion"]),
                     score3(s["nature_followup"])))
    cap = ("Table X.4: Per-strand Nature results — Carajás (inside a KBA: "
           "98.8 % overlap → overlap branch) vs Morowali IMIP (KBA ~6 km "
           "away → distance-decay branch; highest on-site habitat "
           "conversion). Proximity = max(overlap_pct, exp(−dist/10)).")
    _append_table(doc, headers, rows, {1, 2, 4, 5}, cap)
    return "full", None


def appendix_x5(doc, summary):
    order = ["Morowali", "Comodoro", "Jamnagar", "Leshan"]
    headers = ["Site", "VIIRS lit-contrast (state · pct)",
               "Habitat spatial-link (state · offset km)", "Wind states"]
    rows, n_dash = [], 0
    for needle in order:
        s = summary[_site_key(summary, needle)]
        lit = _to_float(s["ghg_viirs_lit_contrast_pct"])
        viirs = f"{s['ghg_viirs_attrib_state'] or DASH} · " + \
                (f"{lit * 100:.1f}%" if lit is not None else DASH)
        off = s["nature_spatial_offset_km"]
        hab = f"{s['nature_habitat_attrib_state'] or DASH} · " + \
              (f"{num(off, 2)} km" if _to_float(off) is not None else DASH)
        wind = (s["wind_attrib_states"] or DASH).replace(";", ", ")
        if _to_float(off) is None:
            n_dash += 1
        rows.append((_short(s["site"]), viirs, hab, wind))
    cap = ("Table X.5: Attributability labels at the flaring sites and the "
           "polysilicon site — VIIRS lit-contrast state (ring percentile), "
           "habitat spatial-link state (change-centroid offset from the "
           "supplier), and per-pollutant wind states. 'sparse' offsets show as "
           "— (too few change pixels to localise).")
    _append_table(doc, headers, rows, set(), cap)
    return ("full" if n_dash == 0 else "partial",
            f"{n_dash} habitat-offset — (sparse)" if n_dash else None)


def appendix_x6(doc, summary):
    a = summary[_site_key(summary, "Norilsk")]
    b = summary[_site_key(summary, "Escondida")]
    headers = ["Site", "Dropped indicators", "Fallbacks fired", "Survivors (Air)",
               "Composite conf."]
    rows = []
    for s in (a, b):
        dropped = (s["dropped_indicators"] or "").split(";") if s["dropped_indicators"] else []
        air_dropped = sum(1 for d in dropped if d.startswith("air."))
        rows.append((_short(s["site"]), _pretty_list(s["dropped_indicators"]),
                     _pretty_list(s["fallbacks_fired"], suffix=" (SPPY)"),
                     f"{9 - air_dropped} of 9", score3(s["composite_confidence"])))
    cap = ("Table X.6: Degradation events (full) — Norilsk vs Escondida. "
           "PM₂.₅ and PM₁₀ drop at every site (5 km buffer < "
           "44.5 km native pixel), renormalising Air to 7 of 9 survivors. "
           "Norilsk additionally recovered SO₂ and AOD via the SPPY "
           "temporal fallback; Escondida fired none.")
    _append_table(doc, headers, rows, {3, 4}, cap)
    return "full", None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    summary = brt._load_summary()
    long = brt._load_long()
    doc = docx.Document(str(SRC_DOCX))
    results = {}

    # --- Figure 7.A replaces the Table 7.4 placeholder (§7.3) ---
    cap_a = ("Figure 7.A: the borrowed combustion proxy is blind to flaring, "
             "which VIIRS catches; where both grammars agree (Escondida, "
             "Carajás) the GHG signal is genuinely low. Divergence is itself the "
             "signal. Source: Appendix X sweep.")
    _replace_placeholder_with_figure(doc, "Table 7.4", FIG_A, cap_a)
    results["Figure 7.A"] = ("rendered", "at §7.3 (Table 7.4 slot)")

    # --- Figure 7.B at §7.5; Table 7.0 placeholder -> cross-reference ---
    cap_b = ("Figure 7.B: two suppliers with opposite risk signatures — an "
             "oil-and-gas field led by its GHG pillar, an iron-ore mine led by "
             "Nature — showing the pillars discriminate between "
             "fundamentally different harms. Source: Appendix X sweep.")
    anchor_b = _find_heading(doc, "Graceful degradation")  # start of §7.6
    _insert_figure_before(doc, anchor_b, FIG_B, cap_b)
    _crossref(doc, "Table 7.0",
              "Figure 7.B (§7.5, “Combining pillars: the composite”) "
              "presents the headline pillar signatures for two contrasting "
              "worked-example suppliers.")
    results["Figure 7.B"] = ("rendered", "at §7.5; 7.0 slot -> cross-ref")

    # --- Merged inline Table 7.7 (confidence + degradation) ---
    insert_merged_77(doc, summary)
    results["Table 7.7 (merged inline)"] = ("full", "confidence + degradation")

    # --- Cross-references for moved tables ---
    xrefs = {
        "Table 7.1": "Full per-indicator results are given in Appendix X, "
                     "Table X.1 (the repeatable core on a single indicator).",
        "Table 7.3": "Full per-indicator Air results are given in Appendix X, "
                     "Table X.2.",
        "Table 7.5": "Full GHG reference/context terms are given in Appendix X, "
                     "Table X.3.",
        "Table 7.6": "Full per-strand Nature results are given in Appendix X, "
                     "Table X.4.",
        "Table 7.10": "Full attributability results are given in Appendix X, "
                      "Table X.5.",
        "Table 7.11": "Full degradation results are given in Appendix X, "
                      "Table X.6.",
    }
    for label, text in xrefs.items():
        _crossref(doc, label, text)
        results[f"{label} -> cross-ref"] = ("moved", "to Appendix X")

    # --- Appendix X at the end ---
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Appendix X: Worked-example screening data", level=2)
    doc.add_paragraph(
        "The tables below give the full per-indicator screening detail "
        "underlying Section 7. All values come from the single-supplier "
        "worked-example sweep (5 km site buffer, engine-default 5×-capped "
        "background ring, latest-valid 90-day window, SPPY temporal fallback "
        "on, all 19 indicators)."
    )
    app_builders = [
        ("X.1", lambda: appendix_x1(doc, long)),
        ("X.2", lambda: appendix_x2(doc, long)),
        ("X.3", lambda: appendix_x3(doc, long)),
        ("X.4", lambda: appendix_x4(doc, summary)),
        ("X.5", lambda: appendix_x5(doc, summary)),
        ("X.6", lambda: appendix_x6(doc, summary)),
    ]
    for label, fn in app_builders:
        status, note = fn()
        results[f"Appendix {label}"] = (status, note)

    doc.save(str(OUT_DOCX))
    print(f"[v2] saved {OUT_DOCX}\n")
    print("=== build summary ===")
    for k, (status, note) in results.items():
        print(f"  {k:30} {status:9} {note or ''}")


if __name__ == "__main__":
    main()
