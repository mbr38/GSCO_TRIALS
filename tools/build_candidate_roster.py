"""Build a standalone, report-styled docx table of the 12 swept candidate sites.

Read-only on the sweep CSV; writes a single-table docx (Table Grid, grey header,
italic caption) that can be dropped into the report. Does not touch the report.

    python tools/build_candidate_roster.py
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import docx

import tools.build_report_tables as brt

REPO = Path(__file__).resolve().parents[1]
CSV = REPO / "tools" / "report_example_sweep.csv"
OUT = REPO / "docs" / "candidate_roster.docx"


def _composite_band(r):
    comp = r["composite_overall"]
    if not comp or r["status"] != "ok":
        return "— (failed)"
    return f"{float(comp):.3f} ({r['composite_band']})"


def _sort_key(r):
    comp = r["composite_overall"]
    # Failed sites (no composite) sink to the bottom.
    return float(comp) if (comp and r["status"] == "ok") else -1.0


def main():
    rows = list(csv.DictReader(CSV.open(encoding="utf-8")))
    rows.sort(key=_sort_key, reverse=True)

    doc = docx.Document()
    headers = ["No.", "Site", "Country / region", "Industry", "Composite (band)"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    right = {0, 4}
    for c, h in enumerate(headers):
        brt._set_cell(table.rows[0].cells[c], h, bold=True, fill=brt.HEADER_FILL,
                      align="right" if c in right else None)
    for i, r in enumerate(rows):
        cells = [str(i + 1), r["site"], r["country"], r["industry"],
                 _composite_band(r)]
        for c, val in enumerate(cells):
            brt._set_cell(table.rows[i + 1].cells[c], val,
                          align="right" if c in right else None)

    cap = doc.add_paragraph(style="Caption")
    cap.add_run(
        "Table X.0: The twelve supplier facilities screened in the worked-example "
        "sweep, spanning GSCO's critical industries (minerals, EV batteries, "
        "solar PV, semiconductors, refining, oil & gas). All screened with the "
        "standard single-supplier defaults (5 km site buffer, engine-default "
        "5×-capped background ring, latest-valid 90-day window, all 19 "
        "indicators). Composite = the equal-weighted mean of the three pillar "
        "follow-up priorities, banded per Appendix C (Low / Moderate / High). "
        "Grasberg returned an Earth Engine error (Dynamic World had no bands in "
        "the persistently cloud-covered Papua highlands) and produced no "
        "composite. Source: Appendix X sweep."
    ).italic = True

    doc.save(str(OUT))
    print(f"wrote {OUT}")
    # Echo the content for the chat.
    print("\n| No. | Site | Country | Industry | Composite (band) |")
    print("|---|---|---|---|---|")
    for i, r in enumerate(rows):
        print(f"| {i+1} | {r['site']} | {r['country']} | {r['industry']} | {_composite_band(r)} |")


if __name__ == "__main__":
    main()
